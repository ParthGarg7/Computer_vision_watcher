#!/usr/bin/env python3
"""
scripts/generate_layer_docs.py
─────────────────────────────────────────────────────────────────────────────
Loop 5 — Generate 5 implementation-based documentation .docx files.

These documents are labelled "BASED ON ACTUAL IMPLEMENTATION" and describe
the code as it IS, not as it was speculated to be in the architecture docs.

Documents generated:
    docs/impl/Layer3_Detection_Implementation_Doc.docx
    docs/impl/Layer4_Identity_Implementation_Doc.docx
    docs/impl/Layer5_Expression_Implementation_Doc.docx
    docs/impl/Registration_CLI_Implementation_Doc.docx
    docs/impl/DeepSORT_ArcFace_Swap_Doc.docx

Run from project root:
    python scripts/generate_layer_docs.py
"""

import os
import sys

# Ensure project root is in sys.path
_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if _ROOT not in sys.path:
    sys.path.insert(0, _ROOT)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = os.path.join(_ROOT, "docs", "impl")

# ─── Styling helpers ─────────────────────────────────────────────────────────

def new_doc(title: str) -> Document:
    """Create a new Document with standard styles."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    h = doc.add_heading(title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x55)
        run.font.size = Pt(18)

    # Label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("⚠ BASED ON ACTUAL IMPLEMENTATION — Not a suggested architecture doc.")
    r.font.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)

    doc.add_paragraph()
    return doc


def add_h1(doc: Document, text: str):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x47, 0x8A)


def add_h2(doc: Document, text: str):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def add_p(doc: Document, text: str, bold: bool = False, italic: bool = False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.bold = bold
    r.font.italic = italic


def add_code(doc: Document, text: str):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x70, 0x1E)
    p.paragraph_format.left_indent = Cm(1)


def add_table(doc: Document, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[1 + r_idx].cells
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Cm(0.5 * (level + 1))


# ─── Document 1: Layer 3 ─────────────────────────────────────────────────────

def build_layer3_doc():
    doc = new_doc("Layer 3 — Face Detection\nImplementation Documentation")

    add_h1(doc, "PART 1 — Technical Deep Dive")

    add_h2(doc, "1.1 What This Layer Does")
    add_p(doc, (
        "Layer 3 (FaceDetector) runs YOLOv8n-face inference on the 640x640 RGB "
        "tensor produced by Layer 2, scales the resulting bounding box coordinates "
        "back to the original frame's pixel space, crops the face regions with a "
        "15% padding margin, and populates ctx.detections with a list of Detection "
        "dataclass instances."
    ))

    add_h2(doc, "1.2 Exact Input / Output Contract")
    add_p(doc, "Input:", bold=True)
    add_code(doc, "ctx.preprocessed_frame  — (640, 640, 3) RGB uint8 for model inference")
    add_code(doc, "ctx.original_frame      — raw BGR for crop extraction")
    add_code(doc, "ctx.original_shape      — (H, W) for coordinate scaling")
    add_p(doc, "Output (per Detection):", bold=True)
    add_code(doc, "bbox_original           — [x1, y1, x2, y2] in original pixel coords")
    add_code(doc, "bbox_resized            — [x1, y1, x2, y2] in 640x640 model space")
    add_code(doc, "confidence              — float 0.0-1.0")
    add_code(doc, "landmarks_original      — [(lx, ly, lconf), ...] 5-point or None")
    add_code(doc, "face_crop               — BGR crop of original_frame with 15% padding")
    add_code(doc, "face_crop_shape         — (H, W) of the crop")

    add_h2(doc, "1.3 Full Tech Stack (Actual Code)")
    add_table(doc,
        ["Component", "Version", "Module", "Why Chosen"],
        [
            ["YOLOv8n-face", "8.4+", "ultralytics.YOLO", "YOLOv8n-face checkpoint (arnabdhar) tuned for facial regions; smallest YOLO variant, real-time capable."],
            ["OpenCV", "4.13", "cv2", "BGR crop extraction and coordinate clamping."],
            ["NumPy", "2.0+", "numpy", "Array operations for coordinate scaling."],
            ["Torch", "2.7+ cu128", "torch", "YOLOv8 inference backend with CUDA 12.8 support."],
        ]
    )

    add_h2(doc, "1.4 Alternatives and When to Use Them")
    add_table(doc,
        ["Alternative", "When to Prefer It"],
        [
            ["RetinaFace (InsightFace)", "Provides 5-point landmarks natively. Prefer when Layer 4 needs precise alignment without SCRFD fallback."],
            ["MediaPipe Face Detection", "Faster on CPU (TFLite backend). Prefer for embedded devices without GPU."],
            ["MTCNN", "Older but well-tested for academic benchmarks. Prefer when reproducibility with prior work is required."],
        ]
    )

    add_h2(doc, "1.5 Key Function/Module Names (Actual Code)")
    add_code(doc, "src/layer3_detection/detector.py — FaceDetector class")
    add_code(doc, "FaceDetector.detect(ctx: FrameContext) -> FrameContext")
    add_code(doc, "FaceDetector._scale_box(box, sx, sy) -> list")
    add_code(doc, "FaceDetector._crop_face(frame, bbox) -> np.ndarray")
    add_code(doc, "src/layer3_detection/validator.py — ValidationPipeline class")

    add_h2(doc, "1.6 Code Style Conventions")
    for item in [
        "Class names: PascalCase (FaceDetector, ValidationPipeline)",
        "Constants: UPPER_SNAKE_CASE (DEFAULT_MODEL_PATH, CROP_PADDING)",
        "Module docstrings explain the layer contract in full before any code",
        "Each step in a pipeline method is marked with a comment (Step 1, Step 2, ...)",
        "All drawing happens on annotated = frame.copy() — original_frame is never mutated",
    ]:
        add_bullet(doc, item)

    add_h1(doc, "PART 2 — Non-Technical Explainer")

    add_h2(doc, "2.1 What Layer 3 Accomplishes")
    add_p(doc, (
        "Layer 3 is the 'eyes' of the system. It receives a prepared video frame "
        "from Layer 2 and asks the AI model: 'where are the faces in this image?' "
        "For each face found, it draws an invisible box around it, records how "
        "confident the model is, and crops out just the face region to pass along "
        "to the next layer. If there are no faces, it reports zero detections "
        "without crashing."
    ))

    add_h2(doc, "2.2 Why It Matters")
    add_p(doc, (
        "Without Layer 3, the pipeline doesn't know there are any people in the "
        "scene. Every downstream layer (identity, expression, analytics) depends "
        "on the face locations and crops this layer provides. The accuracy of "
        "bounding box coordinates here directly determines how well the rest of "
        "the pipeline works."
    ))

    add_h2(doc, "2.3 Prerequisite Concepts for a Beginner")
    for item in [
        "What is a bounding box? — A rectangle defined by its top-left and bottom-right corners, used to mark where an object is in an image.",
        "What is a neural network? — A mathematical function that learns patterns from many examples. YOLOv8n-face learned what faces look like from millions of photos.",
        "What is a confidence score? — A number between 0 and 1 (where 1 is 'completely certain') that the model gives to say how sure it is about each detection.",
        "What is BGR vs RGB? — Computers can store colour in different orders. OpenCV uses Blue-Green-Red; YOLOv8 expects Red-Green-Blue. Layer 2 handles the conversion.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, "Layer3_Detection_Implementation_Doc.docx")
    doc.save(path)
    print(f"  ✅ Saved: {path}")


# ─── Document 2: Layer 4 ─────────────────────────────────────────────────────

def build_layer4_doc():
    doc = new_doc("Layer 4 — Identity\nImplementation Documentation")

    add_h1(doc, "PART 1 — Technical Deep Dive")

    add_h2(doc, "1.1 What This Layer Does")
    add_p(doc, (
        "Layer 4 (FaceIdentifier) performs three tasks in sequence per frame: "
        "(1) extracts a 512-dimensional ArcFace embedding from each face crop via "
        "InsightFace buffalo_l; (2) searches a FAISS IndexFlatIP registry for the "
        "nearest known person and assigns an identity label + similarity score; "
        "(3) updates DeepSORT multi-object tracker with the bboxes and ArcFace "
        "appearance features to assign stable track_id integers across frames."
    ))

    add_h2(doc, "1.2 Exact Input / Output Contract")
    add_p(doc, "Input:", bold=True)
    add_code(doc, "ctx.detections — populated list of Detection objects from Layer 3")
    add_code(doc, "ctx.original_frame — BGR frame passed to DeepSORT.update_tracks()")
    add_p(doc, "Output (added to each Detection):", bold=True)
    add_code(doc, "track_id            — int (DeepSORT track ID) or None if unconfirmed")
    add_code(doc, "identity_label      — str name or 'unknown'")
    add_code(doc, "embedding           — np.ndarray (512,) float32 L2-normalised")
    add_code(doc, "similarity_score    — float cosine similarity 0.0-1.0")
    add_code(doc, "is_known            — bool (score >= threshold 0.45)")
    add_code(doc, "aligned_face        — np.ndarray (112, 112, 3) BGR")

    add_h2(doc, "1.3 Full Tech Stack (Actual Code)")
    add_table(doc,
        ["Component", "Version", "Module", "Why Chosen"],
        [
            ["InsightFace buffalo_l", "1.0.1", "insightface.app.FaceAnalysis", "Bundles SCRFD detector + ArcFace ResNet-50 in one package. normed_embedding is already L2-normalised. 99.83% LFW accuracy."],
            ["FAISS IndexFlatIP", "1.14.3", "faiss", "Exact brute-force inner product search. Correct for cosine similarity on L2-normalised vectors. No training step. MVP scale (<10k people)."],
            ["DeepSORT", "1.3.2", "deep_sort_realtime", "Kalman filter + IoU + appearance embedding matching. Stable track IDs across brief occlusions."],
            ["NumPy", "2.0+", "numpy", "Embedding manipulation, normalisation, FAISS batch operations."],
        ]
    )

    add_h2(doc, "1.4 DeepSORT / ArcFace Appearance Swap")
    add_p(doc, (
        "DeepSORT's default appearance model is a ResNet-50 trained on Market-1501 "
        "pedestrian re-identification data. This was replaced with the ArcFace 512-d "
        "embedding from InsightFace. The swap is zero extra cost (ArcFace is already "
        "computed for recognition). The max_cosine_distance parameter was re-tuned "
        "from 0.2 (pedestrian) to 0.45 (ArcFace faces). Full explanation is in the "
        "tracker.py module docstring."
    ))

    add_h2(doc, "1.5 Alternatives and When to Use Them")
    add_table(doc,
        ["Alternative", "When to Prefer It"],
        [
            ["FAISS IndexHNSWFlat", "When CRUD (delete individual embeddings) is needed — HNSW supports incremental updates. MVP uses IndexFlatIP (rebuild on delete)."],
            ["antelopev2 model pack", "Upgrade to ResNet-100 ArcFace for higher accuracy at the cost of ~30% more inference time."],
            ["ByteTrack / SORT", "If identity appearance matching is not needed — SORT uses IoU only, faster but loses tracks on occlusion faster."],
        ]
    )

    add_h2(doc, "1.6 Key Function/Module Names (Actual Code)")
    add_code(doc, "src/layer4_identity/identifier.py — FaceIdentifier.identify(ctx)")
    add_code(doc, "src/layer4_identity/embedder.py — FaceEmbedder.get_embedding(face_crop)")
    add_code(doc, "src/layer4_identity/identity_store.py — IdentityStore.register(), .search(), .delete_person()")
    add_code(doc, "src/layer4_identity/tracker.py — FaceTracker.update(raw_detections, frame)")
    add_code(doc, "src/layer4_identity/validator.py — Layer4ValidationPipeline")
    add_code(doc, "src/layer4_identity/embedding_validator.py — run_embedding_unit_tests()")

    add_h1(doc, "PART 2 — Non-Technical Explainer")

    add_h2(doc, "2.1 What Layer 4 Accomplishes")
    add_p(doc, (
        "Layer 4 is the 'recognition memory' of the system. When a face is detected, "
        "Layer 4 converts it into a unique numerical signature (called an embedding) "
        "and checks it against a database of registered people. If it finds a match, "
        "it labels the face with that person's name. It also keeps track of where "
        "each person is across multiple frames — so even if someone turns away "
        "briefly, the system still knows it's the same person when they turn back."
    ))

    add_h2(doc, "2.2 Why It Matters")
    add_p(doc, (
        "Without Layer 4, the pipeline can detect faces but not know whose they are. "
        "Identity is the bridge between raw detection and meaningful analytics. "
        "It enables features like 'how long did Alice spend in the meeting room?' "
        "or 'who was present during the presentation?'"
    ))

    add_h2(doc, "2.3 Prerequisite Concepts for a Beginner")
    for item in [
        "What is a face embedding? — A list of 512 numbers that represents 'what this person's face looks like' in a way that can be mathematically compared. Similar faces produce similar numbers.",
        "What is cosine similarity? — A way to measure how similar two embedding vectors are. A score of 1.0 means 'identical face'. A score near 0.0 means 'completely different person'.",
        "What is a tracking algorithm? — Software that follows an object through multiple video frames, giving it a stable ID even when it temporarily disappears.",
        "What is a vector database? — A database optimised for storing and searching these number-lists (embeddings) rather than text or numbers.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, "Layer4_Identity_Implementation_Doc.docx")
    doc.save(path)
    print(f"  ✅ Saved: {path}")


# ─── Document 3: Layer 5 ─────────────────────────────────────────────────────

def build_layer5_doc():
    doc = new_doc("Layer 5 — Expression Analysis\nImplementation Documentation")

    add_h1(doc, "PART 1 — Technical Deep Dive")

    add_h2(doc, "1.1 What This Layer Does")
    add_p(doc, (
        "Layer 5 (ExpressionAnalyser) estimates the facial expression of each "
        "detected person by running a pre-trained ONNX model on their face crop. "
        "It outputs a probability distribution across 8 AffectNet emotion classes "
        "(neutral, happy, sad, surprise, fear, disgust, anger, contempt), the "
        "dominant class, and its confidence score. Inference is throttled to once "
        "every 5 frames per track_id to control CPU usage, with sticky carry-forward "
        "on throttled frames."
    ))

    add_h2(doc, "1.2 Why hsemotion-onnx Instead of DeepFace")
    add_p(doc, (
        "The architecture doc specified DeepFace. However, DeepFace requires "
        "TensorFlow as its backend. TensorFlow does not publish Python 3.14 wheels. "
        "The project venv uses Python 3.14.3. hsemotion-onnx was chosen as the "
        "highest-quality Python 3.14-compatible alternative: it uses ONNX Runtime "
        "(which has 3.14 wheels), is trained on AffectNet (~460K labelled face images), "
        "and provides 8-class output with softmax probabilities."
    ))

    add_h2(doc, "1.3 Exact Input / Output Contract")
    add_p(doc, "Input:", bold=True)
    add_code(doc, "det.face_crop   — BGR uint8 (H, W, 3) from Layer 3")
    add_code(doc, "det.track_id    — int from Layer 4 (used for throttle gating)")
    add_p(doc, "Output (added to Detection):", bold=True)
    add_code(doc, "expression_scores      — dict: {'happy': 0.72, 'neutral': 0.18, ...}")
    add_code(doc, "dominant_expression    — str e.g. 'happy'")
    add_code(doc, "expression_confidence  — float 0.0-1.0")

    add_h2(doc, "1.4 Full Tech Stack (Actual Code)")
    add_table(doc,
        ["Component", "Version", "Module", "Why Chosen"],
        [
            ["hsemotion-onnx (EfficientNet-B0)", "latest", "hsemotion_onnx.facial_emotions.HSEmotionRecognizer", "AffectNet-trained, 8-class, Python 3.14 compatible, ONNX Runtime backend."],
            ["ONNX Runtime", "1.27.0", "onnxruntime", "Inference backend with Python 3.14 wheel support. GPU provider available."],
            ["NumPy", "2.0+", "numpy", "Probability normalisation, argmax."],
        ]
    )

    add_h2(doc, "1.5 Throttling Architecture")
    add_p(doc, (
        "ExpressionAnalyser._frame_counter maps each track_id to a frame count. "
        "Inference runs when count % every_n_frames == 0 (default n=5). "
        "Between analyses, the last result is carried forward from _last_known. "
        "Tracks that leave frame are cleaned up via clear_stale_tracks()."
    ))
    add_code(doc, "EXPRESSION_EVERY_N_FRAMES = 5  # adjustable via constructor")
    add_code(doc, "ExpressionAnalyser.analyse(ctx) — per-frame entry point")
    add_code(doc, "ExpressionAnalyser.clear_stale_tracks(active_track_ids)")

    add_h2(doc, "1.6 Alternatives and When to Use Them")
    add_table(doc,
        ["Alternative", "When to Prefer It"],
        [
            ["DeepFace (TF backend)", "When TensorFlow is available (Python ≤ 3.12). Provides age/gender/race in addition to emotion."],
            ["OpenFace (AU-based)", "When AU (Action Unit) granularity is required for detailed emotion research rather than coarse class labels."],
            ["Custom AffectNet finetune", "When domain-specific accuracy is critical (e.g. fine-grained stress detection). Requires training infrastructure."],
        ]
    )

    add_h2(doc, "1.7 Key Function/Module Names (Actual Code)")
    add_code(doc, "src/layer5_expression/analyser.py — ExpressionAnalyser class")
    add_code(doc, "ExpressionAnalyser.analyse(ctx: FrameContext) -> FrameContext")
    add_code(doc, "ExpressionAnalyser._run_inference(face_crop) -> (dict, str, float)")
    add_code(doc, "ExpressionAnalyser.clear_stale_tracks(active_track_ids)")
    add_code(doc, "src/layer5_expression/validator.py — Layer5ValidationPipeline")

    add_h1(doc, "PART 2 — Non-Technical Explainer")

    add_h2(doc, "2.1 What Layer 5 Accomplishes")
    add_p(doc, (
        "Layer 5 reads a person's facial expression from their face crop. It produces "
        "a list of 8 emotions (happy, sad, angry, neutral, etc.) each with a "
        "probability score saying how likely it is that the person is showing that "
        "expression right now. The dominant emotion — the one with the highest score "
        "— is shown on screen. To avoid overloading the computer, expression analysis "
        "only runs every 5th frame per person; the label stays stable between runs."
    ))

    add_h2(doc, "2.2 Why It Matters")
    add_p(doc, (
        "Expression data enables the business intelligence use cases: measuring "
        "audience enjoyment, tracking interview stress, or monitoring presenter "
        "engagement. Without this layer, the pipeline can tell you who is there "
        "but not how they feel."
    ))

    add_h2(doc, "2.3 Prerequisite Concepts for a Beginner")
    for item in [
        "What is an ONNX model? — A standardised file format for AI models that can run on many different software backends without being retrained.",
        "What is AffectNet? — A large dataset of ~460,000 human face photos that have been labelled with the emotion visible in each photo. Training on this teaches the model to recognise emotions.",
        "What is a probability distribution? — A set of numbers (one per emotion class) that all add up to 100%, each saying how likely that emotion is.",
        "What is frame throttling? — Deliberately skipping some frames to reduce CPU usage. The last-computed result is shown on screen until the next analysis runs.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, "Layer5_Expression_Implementation_Doc.docx")
    doc.save(path)
    print(f"  ✅ Saved: {path}")


# ─── Document 4: Registration CLI ────────────────────────────────────────────

def build_registration_cli_doc():
    doc = new_doc("Face Registration CLI\nImplementation Documentation")

    add_h1(doc, "PART 1 — Technical Deep Dive")

    add_h2(doc, "1.1 What This Tool Does")
    add_p(doc, (
        "scripts/register_face.py is a command-line interface for managing the FAISS "
        "identity registry. It allows operators to: register new people, add multiple "
        "face samples per person, attach/update optional display names, list/search "
        "the registry, and delete people with cascaded embedding removal."
    ))

    add_h2(doc, "1.2 Identity Rules Enforced in Code")
    add_table(doc,
        ["Rule", "How Enforced"],
        [
            ["One UUID per person, generated once at first registration", "uuid.uuid4() on --add-new. UUID is never reused."],
            ["Multiple samples under same UUID (not multiple UUIDs)", "--add-sample requires --person-id; raises ValueError if UUID not found."],
            ["Name is metadata, never identity key", "IdentityStore.search() keys on FAISS row index → UUID, never on name string."],
            ["Two people can share a name and still be distinct", "Each --add-new generates a new UUID regardless of name."],
            ["Delete cascades all embeddings", "IdentityStore.delete_person() rebuilds FAISS index minus deleted rows."],
        ]
    )

    add_h2(doc, "1.3 CLI Modes (Actual Commands)")
    add_code(doc, "python scripts/register_face.py --add-new --image face.jpg --name Alice")
    add_code(doc, "python scripts/register_face.py --add-sample --person-id <UUID> --image face2.jpg")
    add_code(doc, "python scripts/register_face.py --set-name --person-id <UUID> --name Bob")
    add_code(doc, "python scripts/register_face.py --list")
    add_code(doc, "python scripts/register_face.py --search Alice")
    add_code(doc, "python scripts/register_face.py --info --person-id <UUID>")
    add_code(doc, "python scripts/register_face.py --delete --person-id <UUID>")
    add_code(doc, "python scripts/register_face.py --clear")

    add_h2(doc, "1.4 Storage Schema")
    add_table(doc,
        ["File", "Format", "Contents"],
        [
            ["models/identity_store.faiss", "FAISS binary", "L2-normalised 512-d ArcFace embeddings. One row per embedding (multiple rows per person)."],
            ["models/identity_store.meta.json", "JSON", "{meta: {uuid: {name, count}}, id_map: [uuid, uuid, ...], threshold: float}"],
        ]
    )

    add_h2(doc, "1.5 Key Function/Module Names (Actual Code)")
    add_code(doc, "scripts/register_face.py — main(), cmd_add_new(), cmd_add_sample(), cmd_set_name()")
    add_code(doc, "src/layer4_identity/identity_store.py — IdentityStore.register(), .search(), .delete_person(), .save()")
    add_code(doc, "src/layer4_identity/embedder.py — FaceEmbedder.get_embedding(face_crop)")

    add_h1(doc, "PART 2 — Non-Technical Explainer")

    add_h2(doc, "2.1 What the Registration CLI Accomplishes")
    add_p(doc, (
        "The registration tool is the way you 'teach' the system who people are. "
        "Before the pipeline can recognise Alice or Bob, someone must run this tool "
        "to capture their face, generate a unique ID for them, and store it. "
        "Once registered, every time that person appears in the camera, the pipeline "
        "will match their face to their stored ID and show their name."
    ))

    add_h2(doc, "2.2 Why Names Are Not Identity Keys")
    add_p(doc, (
        "Names are just labels for convenience. Two different people could be named "
        "'Alice Smith'. The system identifies people by their unique ID (a random "
        "string of letters and numbers) — never by their name. This prevents two "
        "people being accidentally merged just because they share a name."
    ))

    add_h2(doc, "2.3 Prerequisite Concepts for a Beginner")
    for item in [
        "What is a UUID? — A Universally Unique IDentifier — a randomly generated string (e.g. 550e8400-e29b-41d4-a716-446655440000) that is statistically guaranteed to be unique.",
        "What is a primary key? — In a database, the single field that uniquely identifies each record. Here, person_id (UUID) is the primary key — it never changes.",
        "What is cascade delete? — When you delete a person, their face embeddings are also automatically deleted so no orphaned data remains.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, "Registration_CLI_Implementation_Doc.docx")
    doc.save(path)
    print(f"  ✅ Saved: {path}")


# ─── Document 5: DeepSORT/ArcFace Swap ──────────────────────────────────────

def build_deepsort_swap_doc():
    doc = new_doc("DeepSORT + ArcFace Appearance Swap\nImplementation Documentation")

    add_h1(doc, "PART 1 — Technical Deep Dive")

    add_h2(doc, "1.1 What DeepSORT's Default Appearance Model Does")
    add_p(doc, (
        "DeepSORT (Wojke et al., ICIP 2017) extends Simple SORT (Kalman filter + "
        "IoU matching) with a second association step using a deep appearance "
        "descriptor — a ResNet CNN that produces a 128-d embedding per detection. "
        "The default model is trained on Market-1501, a pedestrian re-identification "
        "dataset of 32,668 annotated images of 1,501 identities captured from "
        "university campus cameras. It encodes features useful for distinguishing "
        "people by clothing, body shape, and gait."
    ))

    add_h2(doc, "1.2 Why the Pedestrian Model is Suboptimal for Face Tracking")
    for item in [
        "Wrong feature space: Market-1501 images are full-body crops (~128x64px). The model learned clothing, torso, leg patterns — none of which are present in a face-only crop.",
        "Wrong input region: For face tracking, the 'detection' bbox is a face region. Running a body-trained CNN on a face crop produces features that are not discriminative for faces.",
        "Lower inter-person discriminability: Two people with similar clothing would confuse the pedestrian model; ArcFace is specifically optimised to separate faces in angular space.",
        "Wasted compute: Running a second CNN model (128-d output, lower quality) when a 512-d face-optimised embedding is already computed is strictly wasteful.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "1.3 Why ArcFace is a Valid Zero-Extra-Cost Substitute")
    for item in [
        "Already computed: FaceEmbedder.get_embedding() runs InsightFace's ArcFace backbone for recognition. The 512-d normed_embedding comes at no additional inference cost.",
        "Face-discriminative: ArcFace (Deng et al., CVPR 2019) was trained with additive angular margin loss on 5.8M faces to maximally separate different-person faces. Same-person embeddings cluster tightly; different-person embeddings are well separated.",
        "L2-normalised: InsightFace's normed_embedding is already L2-normalised. DeepSORT's NearestNeighborDistanceMetric uses cosine distance = 1 - inner product, which is exactly correct for L2-normalised vectors.",
    ]:
        add_bullet(doc, item)

    add_h2(doc, "1.4 Trade-offs and Risks")
    add_table(doc,
        ["Risk", "Severity", "Mitigation"],
        [
            ["Dimensionality mismatch (512-d vs 128-d)", "None", "DeepSORT's metric accepts any fixed-dim vector. No code change needed beyond embedder=None."],
            ["Threshold re-tuning required", "Medium", "Default max_cosine_distance=0.2 was for 128-d pedestrian. Tuned to 0.45 for 512-d ArcFace faces. Calibrate on deployment data."],
            ["Missing embedding on undetected frames", "Low", "DeepSORT's Kalman filter handles position prediction via IoU-only matching when embedding is None. Track expiry is correct."],
            ["ArcFace failure on tiny crops (<40px)", "Low", "FaceEmbedder returns None; DeepSORT falls back to IoU matching for that frame."],
        ]
    )

    add_h2(doc, "1.5 Implementation (Actual Code)")
    add_code(doc, "src/layer4_identity/tracker.py — FaceTracker class")
    add_code(doc, "DeepSort(embedder=None, ...)  # disables internal ReID model download")
    add_code(doc, "deepsort_input.append((bbox_ltwh, confidence, 'face', feature))")
    add_code(doc, "# feature = ArcFace 512-d embedding (None falls back to IoU)")
    add_code(doc, "MAX_COSINE_DISTANCE = 0.45  # tuned from 0.2 (pedestrian default)")

    add_h1(doc, "PART 2 — Non-Technical Explainer")

    add_h2(doc, "2.1 The Problem with the Default Model")
    add_p(doc, (
        "The default 'remembers what people look like' model inside DeepSORT was "
        "designed for tracking people by their clothing and body shape — not their "
        "faces. When you point a camera at people's faces, this model is essentially "
        "looking for features it was never taught about. It still works, but less "
        "accurately — it might occasionally think a different person is someone it "
        "was already tracking."
    ))

    add_h2(doc, "2.2 The Swap: ArcFace as Appearance Descriptor")
    add_p(doc, (
        "Instead of running a second 'looks like' model, we reuse the face embedding "
        "that Layer 4 already computed for recognition. This is like using the "
        "fingerprint you already took to also check 'is this the same person I saw "
        "last frame?' — no extra work, and it's more accurate because fingerprints "
        "(face embeddings) are specifically designed to tell people apart."
    ))

    add_h2(doc, "2.3 Prerequisite Concepts for a Beginner")
    for item in [
        "What is multi-object tracking? — Following multiple moving objects (faces) across consecutive video frames, giving each one a stable ID.",
        "What is a Kalman filter? — A mathematical tool that predicts where a moving object will be in the next frame based on its current speed and direction.",
        "What is re-identification? — The task of deciding 'is this person in frame N the same person I saw in frame N-5?' — even if they briefly left the camera's view.",
        "What is cosine similarity? — A measure of how similar two embedding vectors are, based on the angle between them (not their length). 1.0 = same direction = same person.",
    ]:
        add_bullet(doc, item)

    path = os.path.join(OUTPUT_DIR, "DeepSORT_ArcFace_Swap_Implementation_Doc.docx")
    doc.save(path)
    print(f"  ✅ Saved: {path}")


# ─── Entry point ─────────────────────────────────────────────────────────────

def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"\n  Generating 5 implementation-based documentation files...")
    print(f"  Output directory: {OUTPUT_DIR}\n")

    build_layer3_doc()
    build_layer4_doc()
    build_layer5_doc()
    build_registration_cli_doc()
    build_deepsort_swap_doc()

    print(f"\n  ✅ All 5 documents generated in: {OUTPUT_DIR}\n")


if __name__ == "__main__":
    main()
